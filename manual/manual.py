
# -------------------------------------------------------------------------------------------
#
# manual.py : prende gli screenshots nella direcory e ne fa un manuale Word 
#
# -------------------------------------------------------------------------------------------
#
# written by Massimiliano Cosmelli ( @_°° massimiliano.cosmelli@accelleron-industries.com )
#
#                   CopyRight 2026 Accelleron Industries 
#
# -------------------------------------------------------------------------------------------

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

def crea_manuale(percorso_immagini, output_docx):
    # Crea documento
    doc = Document()

    # Margini (puoi modificarli se vuoi più spazio)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Larghezza utile della pagina
    larghezza_pagina = section.page_width
    margini = section.left_margin + section.right_margin
    larghezza_utilizzabile = larghezza_pagina - margini

    # Lista immagini ordinate
    immagini = sorted([
        f for f in os.listdir(percorso_immagini)
        if f.lower().endswith((".png", ".jpg", ".jpeg", ".bmp"))
    ])

    for i, nome_file in enumerate(immagini):
        path_img = os.path.join(percorso_immagini, nome_file)

        # Inserisci immagine
        p_img = doc.add_paragraph()
        run = p_img.add_run()
        run.add_picture(path_img, width=larghezza_utilizzabile)
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spazio + testo descrizione
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run(f"Descrizione passo {i+1}: ______________________________\n\n")
        run_desc.font.size = Pt(12)

        # Spazio extra per scrivere
        doc.add_paragraph("\n\n")

        # Nuova pagina (tranne ultima immagine)
        if i < len(immagini) - 1:
            doc.add_page_break()

    # Salva documento
    doc.save(output_docx)
    print(f"Manuale creato: {output_docx}")


# === UTILIZZO ===
cartella_immagini = r"C:\Users\ITMACOS\OneDrive - Accelleron-Industries\PC e SMARTPHONE\Procedura x Installazione Win11 privato"
file_output = r"C:\Users\ITMACOS\OneDrive - Accelleron-Industries\PC e SMARTPHONE\Procedura x Installazione Win11 privato\manuale_installazione_windows11.docx"

crea_manuale(cartella_immagini, file_output)

