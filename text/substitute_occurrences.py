from docx import Document
import os

def conta_e_sostituisci_docx(nome_file, stringa_cerca, stringa_sostituisci=None):
    doc = Document(nome_file)
    conteggio = 0

    # Funzione per elaborare paragrafi
    def processa_paragrafi(paragrafi):
        nonlocal conteggio
        for par in paragrafi:
            if stringa_cerca in par.text:
                conteggio += par.text.count(stringa_cerca)
                if stringa_sostituisci is not None:
                    nuovo_testo = par.text.replace(stringa_cerca, stringa_sostituisci)
                    # Sostituzione mantenendo la struttura base
                    for run in par.runs:
                        run.text = ""
                    par.runs[0].text = nuovo_testo

    # Paragrafi normali
    processa_paragrafi(doc.paragraphs)

    # Tabelle
    for tabella in doc.tables:
        for riga in tabella.rows:
            for cella in riga.cells:
                processa_paragrafi(cella.paragraphs)

    return doc, conteggio


def main():
    nome_file = input("Inserisci il nome del file .docx (con estensione): ").strip()

    if not os.path.exists(nome_file):
        print("File non trovato.")
        return

    stringa_cerca = input("Inserisci la stringa da cercare: ")

    doc, occorrenze = conta_e_sostituisci_docx(nome_file, stringa_cerca)

    print(f"\nOccorrenze trovate: {occorrenze}")

    if occorrenze == 0:
        return

    risposta = input("Vuoi sostituire la stringa? (s/n): ").lower()

    if risposta == "s":
        stringa_sostituisci = input("Inserisci la nuova stringa: ")

        doc_modificato, _ = conta_e_sostituisci_docx(
            nome_file,
            stringa_cerca,
            stringa_sostituisci
        )

        nuovo_nome = nome_file.replace(".docx", "_modificato.docx")
        doc_modificato.save(nuovo_nome)

        print(f"Documento salvato come: {nuovo_nome}")
    else:
        print("Nessuna modifica effettuata.")


if __name__ == "__main__":
    main()