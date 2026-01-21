#
# merge_all__doc.py : fa il merge dei documenti .doc, .docx, .odt in un unico .docx
#
#	Input supportati:
#       .docx → letto nativamente
#       .odt → letto via odfpy
#       .doc → tentativo di estrazione testo (solo testo, niente formattazione)
#	Output:
#	    un unico file merge.docx
#	    ogni documento preceduto da titolo = nome file (senza estensione)

from pathlib import Path
from docx import Document
from odf import text, teletype
from odf.opendocument import load

def read_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def read_odt(path):
    doc = load(path)
    paras = doc.getElementsByType(text.P)
    return "\n".join(teletype.extractText(p) for p in paras if teletype.extractText(p).strip())

def merge_documents(input_dir, output_file):
    files = sorted(
        f for f in Path(input_dir).iterdir()
        if f.suffix.lower() in {".docx", ".odt"}
    )

    out = Document()

    for f in files:
        out.add_heading(f.stem, level=1)

        content = read_docx(f) if f.suffix.lower() == ".docx" else read_odt(f)

        for line in content.splitlines():
            out.add_paragraph(line)

        out.add_page_break()

    out.save(output_file)
    print(f"Creato: {output_file}")

# ---------------------------------------------------
# USO
# ---------------------------------------------------

if __name__ == "__main__":
    directory_input = r"D:\miei-scritti\docx"
    output = r"D:\miei-scritti\docx\__MERGE_di_tutti_i_DOC.docx"

    merge_documents(directory_input, output)
